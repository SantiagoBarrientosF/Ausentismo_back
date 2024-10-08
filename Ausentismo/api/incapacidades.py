from Ausentismo.models import *
from Ausentismo.api.serializer import *
from django.http import JsonResponse,HttpResponse
from rest_framework.views import APIView
from django.contrib.auth.models import User
from django.http import JsonResponse 
from django.http import HttpResponse
from docx import Document

class Incapacidadesdata(APIView):
#    authentication_classes = [TokenAuthentication]
#    permission_classes = [IsAuthenticated]
 def get(self,request):    
      Valor= Permisos.objects.filter(tipo_permiso="Licencia remunerada") 
      items_list = []
      for item in Valor: 
         item_dict = {
            'Nombre':item.nombre,
            'Cargo':item.cargo, 
            'Campaña':item.campaña, 
            'Fecha_ingreso_empresa':item.fecha_ingreso_empresa, 
            'Fecha_permiso':item.fecha_peticion, 
         }
         items_list.append(item_dict)
      return JsonResponse(items_list, safe=False) 


def export_word(request):
        permisos = Permisos.objects.all()

        document = Document()

        document.add_heading('Lista de Permisos ' ,0)

        table = document.add_table(rows=1, cols=5)

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Nombre'
        hdr_cells[1].text = 'Cargo'
        hdr_cells[2].text = 'Campaña'
        hdr_cells[3].text = 'Fecha Ingreso Empresa'
        hdr_cells[4].text = 'Fecha Permiso'

        for permiso in permisos:
            row_cells = table.add_row().cells
            row_cells[0].text = permiso.nombre
            row_cells[1].text = permiso.cargo
            row_cells[2].text = permiso.campaña
            row_cells[3].text = str(permiso.fecha_ingreso_empresa)
            row_cells[4].text = str(permiso.fecha_peticion)

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="Informes_Incapacaidades_Andes.docx"'

        document.save(response)

        return response
 
